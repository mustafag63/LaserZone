from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

PURPLE = RGBColor(0x4a, 0x00, 0x80)
BLACK  = RGBColor(0x1a, 0x1a, 0x1a)
GRAY   = RGBColor(0x55, 0x55, 0x55)

def add_title(doc, text):
    p = doc.add_heading(text, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.color.rgb = PURPLE
        run.font.size = Pt(20)
    return p

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = PURPLE
    return p

def add_paragraph(doc, text, bold=False, size=10, color=BLACK, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.7)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = BLACK
    return p

def shade_cell(cell, hex_color="F5F5F5"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, color="CCCCCC"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)

def add_card(table_obj, row_idx, col_idx, title, description):
    cell = table_obj.cell(row_idx, col_idx)
    shade_cell(cell, "F0EBF8")
    set_cell_border(cell, "9B59B6")
    cell.width = Cm(8)

    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_before = Pt(4)
    r1 = p1.add_run("Title: ")
    r1.bold = True
    r1.font.color.rgb = PURPLE
    r1.font.size = Pt(10)
    r2 = p1.add_run(title)
    r2.bold = True
    r2.font.size = Pt(10)
    r2.font.color.rgb = PURPLE

    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(4)
    rd = p2.add_run("Description: ")
    rd.bold = True
    rd.font.size = Pt(9)
    rd.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    rb = p2.add_run(description)
    rb.font.size = Pt(9)
    rb.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# ═══════════════════════════════════════════════════════════════════════════════

add_title(doc, "First Iteration Notes")
doc.add_paragraph()

# ── Requirements section ───────────────────────────────────────────────────────
add_heading(doc, "Requirements until 2nd meeting:", level=2)

p = doc.add_paragraph()
r = p.add_run("Runnable system that has:")
r.bold = True
r.font.size = Pt(11)

bullets = [
    "Login & Register pages (username + password)",
    "Dashboard with left sidebar and user profile (avatar generated from username initial)",
    "My Reservations page — dynamic list (localStorage + API fallback)",
    "Make Reservation popup form (name, players ± buttons, calendar, slot locking)",
    "Backend Auth API — register / login / me (JWT)",
    "Reservation Creation API with conflict check",
    "Slot Availability API (single date & date range)",
    "Open Group Reservation API",
    "Database schema: users, reservations, group_reservations, join_requests",
    "44 passing automated tests (Auth: 12 | Slots: 11 | Reservations: 21)",
]
for b in bullets:
    add_bullet(doc, b)

doc.add_paragraph()
add_heading(doc, "Overall Requirements:", level=2)
doc.add_paragraph()

# ── Cards ──────────────────────────────────────────────────────────────────────
cards_left = [
    ("Login",
     "UI that allows the user to enter the app using username and password. JWT token is issued on successful login."),
    ("Sign-up",
     "UI that allows the user to create an account with username and password. Username must be unique (min 3 chars), password min 6 chars."),
    ("Dashboard – Sidebar",
     "After login, left-hand panel shows user avatar (generated from username initial + color gradient), username, role, navigation menu, and logout button."),
    ("My Reservations",
     "Main menu item. When clicked, main content area shows a dynamic list of the user's active reservations: name, date, time, player count, and Locked badge."),
    ("Make Reservation",
     "Button opens a popup form: reservation name (min 2 chars), players − / + buttons (default 3, min 3, max 20), calendar date & time picker. On confirm, the slot is locked."),
]

cards_right = [
    ("Reservation API",
     "POST /api/reservations — creates a reservation with conflict check. Validates future date, operating hours (10:00–21:00), players (3–20). Returns 409 if slot is fully booked."),
    ("Slot Availability API",
     "GET /api/slots/availability — returns hourly slots (10:00–22:00) with booked/available player counts for a given date or date range (max 30 days)."),
    ("Group Reservation API",
     "POST /api/groups — group leader creates an open reservation with a party size. GET /api/groups lists open groups. Group auto-closes when party size is reached."),
    ("Auth API",
     "POST /api/auth/register and /api/auth/login — JWT-based auth. GET /api/auth/me returns current user. All protected routes require Bearer token."),
    ("Database Schema",
     "Tables: users (username, password, role), reservations (name, date, time, players, status), group_reservations (open/closed/cancelled, party_size, current_count), join_requests (pending/approved/rejected)."),
]

rows = max(len(cards_left), len(cards_right))
tbl = doc.add_table(rows=rows, cols=2)
tbl.style = "Table Grid"

# Remove outer table borders
tbl.style = doc.styles["Normal Table"] if "Normal Table" in [s.name for s in doc.styles] else tbl.style

for i in range(rows):
    tbl.columns[0].width = Cm(8.5)
    tbl.columns[1].width = Cm(8.5)
    if i < len(cards_left):
        add_card(tbl, i, 0, *cards_left[i])
    if i < len(cards_right):
        add_card(tbl, i, 1, *cards_right[i])

doc.add_paragraph()
doc.add_paragraph()

# ── Link ──────────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
r1 = p.add_run("Link: ")
r1.bold = True
r1.font.size = Pt(10)
r2 = p.add_run("https://github.com/mustafag63/laserzone")
r2.font.size = Pt(10)
r2.font.color.rgb = RGBColor(0x00, 0x00, 0xCC)

doc.add_paragraph()

# ── Team ──────────────────────────────────────────────────────────────────────
team = [
    "Mustafa Göçmen",
    "Tuna Öcal",
    "Muhammet Gümüş",
    "Begüm Rana Türkoğlu",
    "Eylül Sena Altunsaray",
]
for name in team:
    p = doc.add_paragraph()
    r = p.add_run(name)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = BLACK

OUTPUT = "Iteration_1_Notes.docx"
doc.save(OUTPUT)
print(f"Word document created: {OUTPUT}")
